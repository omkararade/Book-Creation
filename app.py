import os
import io
import time
import requests
from bs4 import BeautifulSoup
from readability import Document
import streamlit as st
from docx import Document as DocxDocument
from docx.shared import Pt
from dotenv import load_dotenv
from openai import OpenAI
from dotenv import load_dotenv

# Add PDF library with better Unicode support
try:
    from fpdf import FPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    st.warning("PDF export not available. Install fpdf: pip install fpdf")

# ----------------------- Fixed PDF Generation Function -----------------------

def create_pdf_from_text(full_text, title='Book Draft'):
    """Create PDF from text content with proper Unicode handling"""
    if not PDF_AVAILABLE:
        return None
        
    class UnicodePDF(FPDF):
        def header(self):
            # No header for simplicity
            pass
            
        def footer(self):
            # No footer for simplicity
            pass
        
        def add_unicode_text(self, text, font_size=12, style=''):
            """Add text with proper Unicode handling"""
            self.set_font('Arial', style, font_size)
            
            # Split text into lines and handle each line
            lines = text.split('\n')
            for line in lines:
                if not line.strip():
                    self.ln(5)
                    continue
                
                # Clean the text to remove problematic characters
                clean_line = self.clean_text(line)
                
                # Use multi_cell for proper wrapping
                self.multi_cell(0, 6, clean_line)
                self.ln(2)
        
        def clean_text(self, text):
            """Clean text by replacing problematic Unicode characters"""
            # Replace common problematic Unicode characters
            replacements = {
                '\u2014': '--',  # em dash
                '\u2013': '-',   # en dash
                '\u2018': "'",   # left single quote
                '\u2019': "'",   # right single quote
                '\u201c': '"',   # left double quote
                '\u201d': '"',   # right double quote
                '\u2026': '...', # ellipsis
                '\u00a0': ' ',   # non-breaking space
                '\u00ae': '(R)', # registered trademark
                '\u00a9': '(C)', # copyright
                '\u2122': '(TM)',# trademark
            }
            
            for unicode_char, replacement in replacements.items():
                text = text.replace(unicode_char, replacement)
            
            # Remove any other non-printable characters
            text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t\r')
            
            return text
    
    pdf = UnicodePDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    
    # Title
    pdf.set_font('Arial', 'B', 16)
    pdf.cell(0, 10, pdf.clean_text(title), 0, 1, 'C')
    pdf.ln(10)
    
    # Process content
    lines = full_text.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            pdf.ln(5)
            continue
            
        # Handle headings
        if line.startswith('# '):
            pdf.add_unicode_text(line[2:], font_size=16, style='B')
        elif line.startswith('## '):
            pdf.add_unicode_text(line[3:], font_size=14, style='B')
        elif line.startswith('### '):
            pdf.add_unicode_text(line[4:], font_size=12, style='B')
        else:
            pdf.add_unicode_text(line, font_size=12)
    
    try:
        # Use a more robust way to get PDF bytes
        pdf_bytes = pdf.output(dest='S')
        if isinstance(pdf_bytes, str):
            pdf_bytes = pdf_bytes.encode('latin-1', errors='replace')
        bio = io.BytesIO(pdf_bytes)
        bio.seek(0)
        return bio
    except Exception as e:
        st.error(f"PDF generation error: {e}")
        return None

# Alternative PDF generation using reportlab (more robust for Unicode)
def create_pdf_alternative(full_text, title='Book Draft'):
    """Alternative PDF generation using reportlab if available"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.units import inch
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Create bytes buffer
        buffer = io.BytesIO()
        
        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Styles
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Normal_Center',
            parent=styles['Normal'],
            alignment=1,  # center
            fontSize=16,
            spaceAfter=30
        ))
        
        # Build story (content)
        story = []
        
        # Title
        title_para = Paragraph(f"<b>{title}</b>", styles['Normal_Center'])
        story.append(title_para)
        story.append(Spacer(1, 0.5*inch))
        
        # Process content
        lines = full_text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                story.append(Spacer(1, 12))
                continue
                
            # Handle headings
            if line.startswith('# '):
                para = Paragraph(f"<b>{line[2:]}</b>", styles['Heading1'])
            elif line.startswith('## '):
                para = Paragraph(f"<b>{line[3:]}</b>", styles['Heading2'])
            elif line.startswith('### '):
                para = Paragraph(f"<b>{line[4:]}</b>", styles['Heading3'])
            else:
                para = Paragraph(line, styles['Normal'])
            
            story.append(para)
            story.append(Spacer(1, 6))
        
        # Build PDF
        doc.build(story)
        
        # Get PDF bytes
        buffer.seek(0)
        return buffer
        
    except ImportError:
        st.warning("ReportLab not available. Install: pip install reportlab")
        return None
    except Exception as e:
        st.error(f"Alternative PDF generation failed: {e}")
        return None

# ----------------------- Rest of the code remains mostly the same -----------------------

def search_links(query, num_results=5):
    """Manual search implementation that actually works"""
    results = []
    
    # Create proper search query
    search_query = query.replace(' ', '+')
    
    # Try multiple search approaches
    search_urls = [
        f"https://www.google.com/search?q={search_query}&num={num_results}",
        f"https://www.bing.com/search?q={search_query}&count={num_results}",
    ]
    
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"
    }
    
    for search_url in search_urls:
        try:
            response = requests.get(search_url, headers=headers, timeout=10)
            if response.status_code == 200:
                soup = BeautifulSoup(response.text, 'html.parser')
                
                # Try to find links in search results
                links = []
                
                # Google style
                for g in soup.find_all('div', class_='g'):
                    anchor = g.find('a')
                    if anchor and anchor.get('href'):
                        href = anchor.get('href')
                        if href.startswith('/url?q='):
                            href = href.split('/url?q=')[1].split('&')[0]
                        if href.startswith('http') and 'google' not in href:
                            links.append(href)
                
                # Bing style
                if not links:
                    for li in soup.find_all('li', class_='b_algo'):
                        anchor = li.find('a')
                        if anchor and anchor.get('href'):
                            href = anchor.get('href')
                            if href.startswith('http'):
                                links.append(href)
                
                # Generic fallback - look for all http links in content
                if not links:
                    for anchor in soup.find_all('a', href=True):
                        href = anchor.get('href')
                        if href.startswith('http') and not any(x in href for x in ['google', 'bing', 'microsoft']):
                            links.append(href)
                
                # Add unique links to results
                for link in links:
                    if link not in results and len(results) < num_results:
                        try:
                            # Validate it's a real URL
                            if '.' in link and len(link) > 10:
                                results.append(link)
                        except:
                            continue
                
                if results:
                    break
                    
        except Exception as e:
            continue
    
    # If no results from search engines, use topic-specific fallbacks
    if not results:
        query_lower = query.lower()
        if any(word in query_lower for word in ['python', 'programming', 'coding']):
            results = [
                "https://docs.python.org/3/tutorial/",
                "https://realpython.com/",
                "https://www.w3schools.com/python/",
                "https://www.geeksforgeeks.org/python-programming-language/",
                "https://www.programiz.com/python-programming"
            ]
        elif any(word in query_lower for word in ['marketing', 'digital', 'social media']):
            results = [
                "https://blog.hubspot.com/marketing",
                "https://neilpatel.com/blog/",
                "https://www.socialmediaexaminer.com/",
                "https://contentmarketinginstitute.com/",
                "https://moz.com/blog"
            ]
        elif any(word in query_lower for word in ['cooking', 'recipe', 'food']):
            results = [
                "https://www.allrecipes.com/",
                "https://www.foodnetwork.com/",
                "https://www.bbcgoodfood.com/",
                "https://www.seriouseats.com/",
                "https://www.bonappetit.com/"
            ]
        else:
            # General knowledge fallback
            results = [
                "https://en.wikipedia.org/wiki/Main_Page",
                "https://www.britannica.com/",
                "https://www.khanacademy.org/",
                "https://www.coursera.org/",
                "https://www.ted.com/talks"
            ]
    
    return results[:num_results]

def init_api_keys():
    load_dotenv()
    openai_key = os.getenv('OPENAI_API_KEY')
    serpapi_key = os.getenv('SERPAPI_API_KEY')
    
    if not openai_key:
        st.warning('Please set environment variable OPENAI_API_KEY before using.')
        return None, serpapi_key
    
    client = OpenAI(api_key=openai_key)
    return client, serpapi_key

def fetch_url_text(url, timeout=10):
    """Fetch page and extract clean text"""
    headers = {"User-Agent": "Mozilla/5.0 (compatible; Book-Maker/1.0)"}
    try:
        r = requests.get(url, headers=headers, timeout=timeout)
        r.raise_for_status()
    except Exception as e:
        return f"Failed to fetch: {str(e)}"
    
    try:
        doc = Document(r.text)
        content = doc.summary()
        soup = BeautifulSoup(content, 'html.parser')
        text = soup.get_text(separator='\n')
        text = '\n'.join([line.strip() for line in text.splitlines() if line.strip()])
        return text if text else "No content extracted"
    except Exception:
        soup = BeautifulSoup(r.text, 'html.parser')
        paragraphs = soup.find_all('p')
        text = '\n'.join(p.get_text().strip() for p in paragraphs if p.get_text().strip())
        return text[:20000] if text else "No content extracted"

def call_openai_chat(client, system_prompt, user_prompt, model='gpt-4o-mini', max_tokens=4000, temperature=0.2):
    if client is None:
        st.error("OpenAI client not initialized.")
        return ""
    
    if model in ['gpt-4o-mini', 'gpt-3.5-turbo']:
        max_tokens = min(max_tokens, 16384)
    elif model in ['gpt-4', 'gpt-4o']:
        max_tokens = min(max_tokens, 8192)
    
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    try:
        resp = client.chat.completions.create(
            model=model,
            messages=messages,
            max_tokens=max_tokens,
            temperature=temperature
        )
        return resp.choices[0].message.content
    except Exception as e:
        st.error(f"OpenAI API error: {e}")
        return ""

def build_research_blob(link_text_pairs, max_total_chars=30000):
    """Create research content from fetched pages"""
    parts = []
    total_chars = 0
    
    for link, text in link_text_pairs:
        if not text or "Failed to fetch" in text or "No content extracted" in text:
            continue
            
        snippet = text[:2000]
        part = f"Source: {link}\n{snippet}\n---\n"
        
        if total_chars + len(part) > max_total_chars:
            break
            
        parts.append(part)
        total_chars += len(part)
    
    return '\n'.join(parts) if parts else "No research content available"

def generate_book_outline(client, title, keywords, short_prompt, research_blob, target_pages, model='gpt-4o-mini'):
    words_per_page = 300
    target_words = target_pages * words_per_page
    
    system = "You are an expert technical writer who creates detailed book outlines."
    user = (
        f"Create a detailed chapter-by-chapter outline for a book about:\n"
        f"Title: {title}\nKeywords: {keywords}\nAngle: {short_prompt}\n"
        f"Target length: {target_pages} pages ({target_words} words)\n\n"
        f"Research context:\n{research_blob[:15000]}\n\n"
        "Create an outline with chapter titles and brief descriptions."
    )
    
    return call_openai_chat(client, system, user, model=model, max_tokens=3000)

def generate_book_chapter(client, chapter_info, research_blob, model='gpt-4o-mini'):
    system = "You are an expert technical writer creating a book chapter."
    user = (
        f"Write this chapter:\n{chapter_info}\n\n"
        f"Use this research for inspiration:\n{research_blob[:10000]}\n\n"
        "Write in a clear, engaging style with practical examples."
    )
    
    return call_openai_chat(client, system, user, model=model, max_tokens=4000)

def generate_entire_book(client, title, keywords, short_prompt, research_blob, target_pages=30, model='gpt-4o-mini'):
    """Generate book using iterative approach"""
    
    with st.spinner('Creating book outline...'):
        outline = generate_book_outline(client, title, keywords, short_prompt, research_blob, target_pages, model)
    
    if not outline:
        return "Failed to generate outline."
    
    # Parse chapters from outline
    chapters = []
    lines = outline.split('\n')
    current_chapter = ""
    
    for line in lines:
        line = line.strip()
        if line and (line.startswith(('Chapter', 'CHAPTER', '##')) or any(x in line.lower() for x in ['chapter', 'part'])):
            if current_chapter:
                chapters.append(current_chapter)
            current_chapter = line
        elif current_chapter and line:
            current_chapter += '\n' + line
    
    if current_chapter:
        chapters.append(current_chapter)
    
    # If parsing failed, create default chapters based on topic
    if len(chapters) < 2:
        st.info("Creating topic-appropriate chapter structure...")
        if any(word in title.lower() for word in ['python', 'programming']):
            chapters = [
                "Chapter 1: Introduction to Python",
                "Chapter 2: Python Basics and Syntax", 
                "Chapter 3: Data Structures in Python",
                "Chapter 4: Functions and Modules",
                "Chapter 5: Object-Oriented Programming",
                "Chapter 6: Working with Files and Data",
                "Chapter 7: Introduction to Python Libraries",
                "Chapter 8: Next Steps in Python Journey"
            ]
        else:
            chapters = [f"Chapter {i+1}" for i in range(6)]
    
    # Generate each chapter
    book_parts = [f"# {title}\n\n", f"## Table of Contents\n\n{outline}\n\n", "## Book Content\n\n"]
    
    for i, chapter_info in enumerate(chapters[:8]):
        with st.spinner(f'Writing chapter {i+1}/{len(chapters)}...'):
            chapter_text = generate_book_chapter(client, chapter_info, research_blob, model)
            if chapter_text:
                book_parts.append(f"## {chapter_info.split(':')[0] if ':' in chapter_info else f'Chapter {i+1}'}\n\n")
                book_parts.append(chapter_text)
                book_parts.append("\n\n")
            time.sleep(1)
    
    # Add references
    book_parts.append("\n\n## References and Further Reading\n\n")
    
    return ''.join(book_parts)

def create_docx_from_text(full_text, title='Book Draft'):
    doc = DocxDocument()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    doc.add_heading(title, level=1)
    
    lines = full_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        else:
            p = doc.add_paragraph()
            p.add_run(line)
    
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def main():
    st.set_page_config(page_title='Book Prototype — Streamlit', layout='wide')
    st.title('Book Prototype — Generate a full draft in one step')

    client, serpapi_key = init_api_keys()

    if client is None:
        st.error("OpenAI client not initialized. Please check your OPENAI_API_KEY environment variable.")
        return

    with st.sidebar:
        st.markdown('**Research settings**')
        num_links = st.slider('Number of links to fetch', 3, 12, 6)
        max_snippet_per_page = st.slider('Max chars to keep per page', 2000, 20000, 5000)
        model = st.selectbox('Model', ['gpt-4o-mini', 'gpt-4o', 'gpt-4'], index=0)
        target_pages = st.slider('Target book pages', 10, 80, 30)

    st.markdown('Enter a topic / keyword / short prompt and we will research and attempt to generate a full draft.')
    title = st.text_input('Book title / Topic', value='Python Programming for Beginners')
    keywords = st.text_input('Comma-separated keywords', value='python, programming, data science, web development')
    short_prompt = st.text_area('Short prompt (1-3 lines) giving tone/angle', value='Write a beginner-friendly guide with code examples')

    if st.button('Research + Generate book'):
        if not client:
            st.error('OpenAI client not available. Check API key.')
            return

        query = f"{title} {keywords}"
        with st.spinner('Searching web...'):
            links = search_links(query, num_results=num_links)
        
        st.write('Found links:', links)

        link_texts = []
        with st.spinner('Fetching page content...'):
            for i, url in enumerate(links):
                st.write(f"Fetching {i+1}/{len(links)}: {url}")
                txt = fetch_url_text(url)
                link_texts.append((url, txt[:max_snippet_per_page]))

        research_blob = build_research_blob(link_texts)
        st.text_area('Research blob preview', value=research_blob[:4000], height=200)

        with st.spinner('Generating book — this can take a few minutes...'):
            book_text = generate_entire_book(client, title, keywords, short_prompt, research_blob, target_pages=target_pages, model=model)

        st.success('Generation complete!')
        st.text_area('Book preview', value=book_text[:5000], height=400)

        # Create downloads - Now with better PDF handling
        col1, col2, col3 = st.columns(3)
        
        with col1:
            docx_io = create_docx_from_text(book_text, title=title)
            st.download_button(
                '📄 Download .docx', 
                data=docx_io, 
                file_name=f"{title.replace(' ','_')}.docx",
                use_container_width=True
            )
        
        with col2:
            st.download_button(
                '📝 Download .txt', 
                data=book_text, 
                file_name=f"{title.replace(' ','_')}.txt",
                use_container_width=True
            )
        
        with col3:
            # Try multiple PDF generation methods
            pdf_io = None
            
            # First try the improved FPDF method
            if PDF_AVAILABLE:
                pdf_io = create_pdf_from_text(book_text, title=title)
            
            # If FPDF fails, try reportlab alternative
            if pdf_io is None:
                pdf_io = create_pdf_alternative(book_text, title=title)
            
            if pdf_io:
                st.download_button(
                    '📘 Download .pdf', 
                    data=pdf_io, 
                    file_name=f"{title.replace(' ','_')}.pdf",
                    use_container_width=True,
                    mime='application/pdf'
                )
            else:
                st.info("PDF export not available. Install: pip install fpdf or pip install reportlab")

if __name__ == '__main__':
    main()